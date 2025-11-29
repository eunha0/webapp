#!/usr/bin/env python3
"""
Convert Word rubric documents to HTML format
"""
import os
from docx import Document

def convert_docx_to_html(docx_path, output_path):
    """Convert DOCX to simple HTML"""
    doc = Document(docx_path)
    
    html_content = []
    html_content.append('<!DOCTYPE html>')
    html_content.append('<html lang="ko">')
    html_content.append('<head>')
    html_content.append('    <meta charset="UTF-8">')
    html_content.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')
    html_content.append('    <style>')
    html_content.append('        body { font-family: "Noto Sans KR", sans-serif; line-height: 1.6; padding: 20px; max-width: 900px; margin: 0 auto; }')
    html_content.append('        h1 { color: #1e3a8a; border-bottom: 3px solid #1e3a8a; padding-bottom: 10px; }')
    html_content.append('        h2 { color: #2563eb; margin-top: 30px; }')
    html_content.append('        table { width: 100%; border-collapse: collapse; margin: 20px 0; }')
    html_content.append('        th, td { border: 1px solid #d1d5db; padding: 12px; text-align: left; }')
    html_content.append('        th { background-color: #f3f4f6; font-weight: bold; color: #1e3a8a; }')
    html_content.append('        tr:nth-child(even) { background-color: #f9fafb; }')
    html_content.append('        p { margin: 10px 0; }')
    html_content.append('        strong { color: #1e40af; }')
    html_content.append('    </style>')
    html_content.append('</head>')
    html_content.append('<body>')
    
    in_table = False
    table_data = []
    
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            # Process table
            table = None
            for tbl in doc.tables:
                if tbl._element == element:
                    table = tbl
                    break
            
            if table:
                html_content.append('    <table>')
                for i, row in enumerate(table.rows):
                    html_content.append('        <tr>')
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if i == 0:
                            html_content.append(f'            <th>{cell_text}</th>')
                        else:
                            html_content.append(f'            <td>{cell_text}</td>')
                    html_content.append('        </tr>')
                html_content.append('    </table>')
        
        elif element.tag.endswith('p'):
            # Process paragraph
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        if para.style.name.startswith('Heading 1') or para.runs and para.runs[0].bold and len(text) < 100:
                            html_content.append(f'    <h1>{text}</h1>')
                        elif para.style.name.startswith('Heading'):
                            html_content.append(f'    <h2>{text}</h2>')
                        else:
                            # Check if text contains bold parts
                            if any(run.bold for run in para.runs):
                                para_html = '    <p>'
                                for run in para.runs:
                                    if run.bold:
                                        para_html += f'<strong>{run.text}</strong>'
                                    else:
                                        para_html += run.text
                                para_html += '</p>'
                                html_content.append(para_html)
                            else:
                                html_content.append(f'    <p>{text}</p>')
                    break
    
    html_content.append('</body>')
    html_content.append('</html>')
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_content))
    
    print(f"Converted: {os.path.basename(docx_path)} -> {os.path.basename(output_path)}")

def main():
    """Main conversion process"""
    source_dir = '/home/user/uploaded_files'
    output_dir = '/home/user/webapp-ai/public/rubric-docs'
    
    # List of files to convert
    files = [
        '뉴욕 주 리젠트 시험 논증적 글쓰기 루브릭.docx',
        '뉴욕 주 리젠트 시험 분석적 글쓰기 루브릭.docx',
        '뉴욕 주 중학교 논술 루브릭.docx',
        '뉴욕 주 초등학교 논술 루브릭.docx',
        'IB 중등 프로그램 고등학교 개인과 사회 논술 루브릭.docx',
        'IB 중등 프로그램 중학교 개인과 사회 논술 루브릭.docx',
        'IB 중등 프로그램 과학 논술 루브릭.docx',
    ]
    
    for filename in files:
        docx_path = os.path.join(source_dir, filename)
        html_filename = filename.replace('.docx', '.html')
        html_path = os.path.join(output_dir, html_filename)
        
        if os.path.exists(docx_path):
            convert_docx_to_html(docx_path, html_path)
        else:
            print(f"File not found: {filename}")
    
    print("\nConversion complete!")

if __name__ == '__main__':
    main()
